from docx import Document
from deep_translator import GoogleTranslator

# Load the document
doc_path = ''
document = Document(doc_path)

# Initialize translator
translator = GoogleTranslator(source='en', target='ro')

# Function to translate text in chunks if it exceeds 5000 characters
def translate_text(text: str):
    max_chars = 5000
    if len(text.strip()) == 0:  # Skip empty strings
        return ""
    elif len(text) <= max_chars:
        return translator.translate(text.strip())
    else:
        # Split text into chunks under 5000 characters, ensuring no empty chunks
        chunks = [text[i:i + max_chars].strip() for i in range(0, len(text), max_chars)]
        translated_chunks = [translator.translate(chunk) for chunk in chunks if len(chunk) > 0]
        return ''.join(translated_chunks)

# Read, translate, and store paragraphs
translated_text = []

# Translate each paragraph in the document
for para in document.paragraphs:
    text = para.text.strip()
    if text:  # Only translate non-empty paragraphs
        translated_text.append(translate_text(text))
    else:
        translated_text.append("")  # Maintain empty lines for formatting

# Create a new document with translated text
translated_doc_path = 'low-histamine-cookbook-ro.docx'
translated_doc = Document()

for para_text in translated_text:
    translated_doc.add_paragraph(para_text)

# Save the translated document
translated_doc.save(translated_doc_path)
